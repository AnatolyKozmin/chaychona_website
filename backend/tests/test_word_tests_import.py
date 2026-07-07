"""Тесты парсера Word для импорта тестов (нужен python-docx)."""

import io

import pytest

docx = pytest.importorskip("docx")


def test_parse_one_question_same_line_options() -> None:
    from docx import Document

    from app.word_tests_import import parse_docx_to_questions

    d = Document()
    p = d.add_paragraph()
    p.add_run("1. Сколько будет 2+2? ")
    p.add_run("A) 3")
    p.add_run(" B) ")
    r = p.add_run("4")
    r.bold = True
    p.add_run(" C) 5")

    buf = io.BytesIO()
    d.save(buf)
    questions, errors = parse_docx_to_questions(buf.getvalue())
    assert errors == []
    assert len(questions) == 1
    assert questions[0]["question_type"] == "single"
    assert questions[0]["text"] == "Сколько будет 2+2?"
    opts = questions[0]["options"]
    assert len(opts) == 3
    assert [o["text"] for o in opts] == ["3", "4", "5"]
    assert [o["is_correct"] for o in opts] == [False, True, False]


def test_parse_multiline_question() -> None:
    from docx import Document

    from app.word_tests_import import parse_docx_to_questions

    d = Document()
    d.add_paragraph("1. Первая строка вопроса")
    d.add_paragraph("вторая строка без номера")
    d.add_paragraph("A) нет")
    p = d.add_paragraph()
    p.add_run("B) ")
    r = p.add_run("да")
    r.bold = True

    buf = io.BytesIO()
    d.save(buf)
    questions, errors = parse_docx_to_questions(buf.getvalue())
    assert errors == []
    assert len(questions) == 1
    assert "Первая строка" in questions[0]["text"]
    assert "вторая строка" in questions[0]["text"]
    assert questions[0]["options"][1]["is_correct"] is True


def _doc_with_good_and_bad_question() -> bytes:
    from docx import Document

    d = Document()
    # хороший вопрос
    p = d.add_paragraph()
    p.add_run("1. Хороший вопрос? A) нет B) ")
    r = p.add_run("да")
    r.bold = True
    # плохой: ни один вариант не выделен жирным
    d.add_paragraph("2. Плохой вопрос? A) раз B) два")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_strict_mode_blocks_whole_file_on_bad_question() -> None:
    from app.word_tests_import import parse_docx_to_questions

    questions, errors = parse_docx_to_questions(_doc_with_good_and_bad_question())
    assert questions == []
    assert len(errors) == 1
    assert "Вопрос 2" in errors[0]


def test_lenient_mode_returns_questions_with_warnings() -> None:
    from app.word_tests_import import parse_docx_to_questions

    questions, warnings = parse_docx_to_questions(_doc_with_good_and_bad_question(), strict=False)
    assert len(questions) == 2
    assert len(warnings) == 1
    assert "Вопрос 2" in warnings[0]
    # проблемный вопрос включён как есть — без отмеченных правильных ответов
    assert all(o["is_correct"] is False for o in questions[1]["options"])


def test_lenient_mode_no_warnings_on_clean_file() -> None:
    from docx import Document

    from app.word_tests_import import parse_docx_to_questions

    d = Document()
    p = d.add_paragraph()
    p.add_run("1. Вопрос? A) нет B) ")
    r = p.add_run("да")
    r.bold = True
    buf = io.BytesIO()
    d.save(buf)

    questions, warnings = parse_docx_to_questions(buf.getvalue(), strict=False)
    assert warnings == []
    assert len(questions) == 1
